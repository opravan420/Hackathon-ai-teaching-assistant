import docx
from .base import BaseDocumentParser
from .exceptions import DocumentExtractionError, NoExtractableTextError, InvalidDocumentError

class DOCXParser(BaseDocumentParser):
    """Parser for extracting text from DOCX files using python-docx."""

    def extract_text(self, file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            raise InvalidDocumentError(f"Corrupted or invalid DOCX file: {str(e)}")

        try:
            extracted = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    extracted.append(para.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        extracted.append(" | ".join(row_text))

            full_text = "\n".join(extracted).strip()
            if not full_text:
                raise NoExtractableTextError("DOCX contains no extractable text.")

            return full_text
        except NoExtractableTextError:
            raise
        except Exception as e:
            raise DocumentExtractionError(f"Error extracting text from DOCX: {str(e)}")
