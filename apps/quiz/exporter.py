import io
import re
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, KeepTogether, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .models import Quiz

def sanitize_filename(name: str) -> str:
    if not name:
        return "AI_Quiz"
    cleaned = re.sub(r'[^\w\s-]', '', name).strip()
    return re.sub(r'[-\s]+', '_', cleaned) or "AI_Quiz"

class QuizExporter:
    """Service to export finalized Quizzes to printable Question Paper and Answer Key documents."""

    @staticmethod
    def generate_question_paper_pdf(quiz: Quiz) -> bytes:
        """Generates Question Paper PDF strictly excluding correct answers and explanations."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle('QP_Title', parent=styles['Heading1'], fontSize=16, leading=20, spaceAfter=6, textColor='#0f172a', alignment=1)
        sub_style = ParagraphStyle('QP_Sub', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=18, textColor='#475569', alignment=1)
        q_style = ParagraphStyle('QP_Q', parent=styles['Heading2'], fontSize=11, leading=15, spaceAfter=6, textColor='#1e293b')
        opt_style = ParagraphStyle('QP_Opt', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=4, textColor='#334155', leftIndent=14)

        topic_display = quiz.topic.strip() if quiz.topic else (quiz.source_file_name or "General Topic")
        story = [
            Paragraph(f"<b>QUESTION PAPER: {topic_display.upper()}</b>", title_style),
            Paragraph(f"Difficulty: {quiz.difficulty.capitalize()} | Total Questions: {quiz.num_questions}", sub_style),
            HRFlowable(width="100%", thickness=1, color="#cbd5e1", spaceAfter=18)
        ]

        labels = ['A', 'B', 'C', 'D']
        for i, q in enumerate(quiz.questions.all(), 1):
            # Escape XML special chars for ReportLab
            q_text_safe = q.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            q_block = [Paragraph(f"<b>{i}. {q_text_safe}</b>", q_style)]
            
            options = list(q.options.all())
            for idx, opt in enumerate(options[:4]):
                lbl = labels[idx] if idx < 4 else str(idx + 1)
                opt_text_safe = opt.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                q_block.append(Paragraph(f"<b>{lbl}.</b> {opt_text_safe}", opt_style))
            
            q_block.append(Spacer(1, 10))
            story.append(KeepTogether(q_block))

        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes

    @staticmethod
    def generate_question_paper_docx(quiz: Quiz) -> bytes:
        """Generates Question Paper DOCX strictly excluding correct answers and explanations."""
        doc = DocxDocument()
        topic_display = quiz.topic.strip() if quiz.topic else (quiz.source_file_name or "General Topic")
        
        # Header
        h = doc.add_heading(f"QUESTION PAPER: {topic_display.upper()}", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_meta = doc.add_paragraph(f"Difficulty: {quiz.difficulty.capitalize()} | Total Questions: {quiz.num_questions}")
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("=" * 60)

        labels = ['A', 'B', 'C', 'D']
        for i, q in enumerate(quiz.questions.all(), 1):
            p_q = doc.add_paragraph()
            r_q = p_q.add_run(f"{i}. {q.text}")
            r_q.bold = True
            
            options = list(q.options.all())
            for idx, opt in enumerate(options[:4]):
                lbl = labels[idx] if idx < 4 else str(idx + 1)
                p_opt = doc.add_paragraph()
                p_opt.paragraph_format.left_indent = Inches(0.25)
                p_opt.add_run(f"{lbl}. ").bold = True
                p_opt.add_run(opt.text)
            
            doc.add_paragraph()

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes

    @staticmethod
    def generate_answer_key_pdf(quiz: Quiz) -> bytes:
        """Generates Answer Key PDF with correct option letters, texts, and explanations."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle('AK_Title', parent=styles['Heading1'], fontSize=16, leading=20, spaceAfter=6, textColor='#065f46', alignment=1)
        sub_style = ParagraphStyle('AK_Sub', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=18, textColor='#047857', alignment=1)
        q_style = ParagraphStyle('AK_Q', parent=styles['Heading2'], fontSize=11, leading=15, spaceAfter=4, textColor='#1e293b')
        ans_style = ParagraphStyle('AK_Ans', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=4, textColor='#047857', leftIndent=14)
        exp_style = ParagraphStyle('AK_Exp', parent=styles['Normal'], fontSize=9, leading=13, spaceAfter=8, textColor='#475569', leftIndent=14)

        topic_display = quiz.topic.strip() if quiz.topic else (quiz.source_file_name or "General Topic")
        story = [
            Paragraph(f"<b>ANSWER KEY: {topic_display.upper()}</b>", title_style),
            Paragraph(f"Difficulty: {quiz.difficulty.capitalize()} | Total Questions: {quiz.num_questions}", sub_style),
            HRFlowable(width="100%", thickness=1, color="#a7f3d0", spaceAfter=18)
        ]

        labels = ['A', 'B', 'C', 'D']
        for i, q in enumerate(quiz.questions.all(), 1):
            q_text_safe = q.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            q_block = [Paragraph(f"<b>{i}. {q_text_safe}</b>", q_style)]
            
            options = list(q.options.all())
            correct_opt = None
            correct_lbl = 'A'
            for idx, opt in enumerate(options[:4]):
                if opt.is_correct:
                    correct_opt = opt
                    correct_lbl = labels[idx] if idx < 4 else str(idx + 1)
                    break

            if correct_opt:
                opt_text_safe = correct_opt.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                q_block.append(Paragraph(f"<b>Correct Answer: {correct_lbl}.</b> {opt_text_safe}", ans_style))
            else:
                q_block.append(Paragraph("<b>Correct Answer:</b> Not specified", ans_style))

            if q.explanation:
                exp_text_safe = q.explanation.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                q_block.append(Paragraph(f"<i>Explanation: {exp_text_safe}</i>", exp_style))

            q_block.append(Spacer(1, 8))
            story.append(KeepTogether(q_block))

        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes

    @staticmethod
    def generate_answer_key_docx(quiz: Quiz) -> bytes:
        """Generates Answer Key DOCX with correct option letters, texts, and explanations."""
        doc = DocxDocument()
        topic_display = quiz.topic.strip() if quiz.topic else (quiz.source_file_name or "General Topic")
        
        # Header
        h = doc.add_heading(f"ANSWER KEY: {topic_display.upper()}", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_meta = doc.add_paragraph(f"Difficulty: {quiz.difficulty.capitalize()} | Total Questions: {quiz.num_questions}")
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("=" * 60)

        labels = ['A', 'B', 'C', 'D']
        for i, q in enumerate(quiz.questions.all(), 1):
            p_q = doc.add_paragraph()
            r_q = p_q.add_run(f"{i}. {q.text}")
            r_q.bold = True

            options = list(q.options.all())
            correct_opt = None
            correct_lbl = 'A'
            for idx, opt in enumerate(options[:4]):
                if opt.is_correct:
                    correct_opt = opt
                    correct_lbl = labels[idx] if idx < 4 else str(idx + 1)
                    break

            p_ans = doc.add_paragraph()
            p_ans.paragraph_format.left_indent = Inches(0.25)
            r_hdr = p_ans.add_run("Correct Answer: ")
            r_hdr.bold = True
            if correct_opt:
                r_val = p_ans.add_run(f"{correct_lbl}. {correct_opt.text}")
                r_val.font.color.rgb = RGBColor(4, 120, 87)
            else:
                p_ans.add_run("Not specified")

            if q.explanation:
                p_exp = doc.add_paragraph()
                p_exp.paragraph_format.left_indent = Inches(0.25)
                r_exp = p_exp.add_run(f"Explanation: {q.explanation}")
                r_exp.italic = True

            doc.add_paragraph()

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes
