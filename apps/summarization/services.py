import io
import logging
from typing import Optional
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from apps.summarization.models import LectureSummary
from apps.ai_engine.document_processing.service import DocumentService
from apps.ai_engine.prompting.prompt_builder import PromptBuilder
from apps.ai_engine.services.llm_service import LLMService

logger = logging.getLogger(__name__)

class SummarizationError(Exception):
    """Raised when lecture summarization generation fails."""
    pass

class SummarizationService:
    """Core service for AI Lecture Note Summarization and Document Export."""

    def __init__(self):
        self.doc_service = DocumentService()
        self.prompt_builder = PromptBuilder()
        self.llm_service = LLMService()

    def generate_summary(
        self,
        teacher,
        uploaded_file,
        custom_instruction: str = ""
    ) -> LectureSummary:
        """
        Generates a structured lecture summary from uploaded document using Gemma 3 4B.
        Note on Document Context Limitation:
        Initial MVP strategy limits context to 4,000 characters (full text for short documents,
        or representative sampled sections across pages/slides for large documents).
        """
        if not uploaded_file:
            raise ValueError("Please upload a lecture document to summarize.")

        # Step 1: Extract document text
        document = self.doc_service.process_document(teacher, uploaded_file)
        extracted_text = document.extracted_text.strip()

        if not extracted_text:
            raise SummarizationError(f"No text could be extracted from '{document.original_filename}'.")

        # Step 2: Context sampling (4000 char MVP strategy)
        if len(extracted_text) <= 4000:
            formatted_context = f"SOURCE: {document.original_filename}\n{extracted_text}"
        else:
            # Sample paragraphs evenly across document to fit context limit
            paragraphs = [p.strip() for p in extracted_text.split("\n\n") if p.strip()]
            sampled = []
            step = max(1, len(paragraphs) // 10)
            for i in range(0, len(paragraphs), step):
                sampled.append(paragraphs[i])
                if sum(len(p) for p in sampled) >= 3800:
                    break
            sampled_text = "\n\n".join(sampled)[:3800]
            formatted_context = f"SOURCE: {document.original_filename} (Sampled Sections)\n{sampled_text}"

        # Step 3: Build prompt & call LLM
        user_prompt, system_prompt = self.prompt_builder.build_summary_prompt(
            formatted_context=formatted_context,
            custom_instruction=custom_instruction
        )

        summary_text = self.llm_service.generate_text(prompt=user_prompt, system_prompt=system_prompt)

        if not summary_text or len(summary_text.strip()) < 20:
            raise SummarizationError("AI generated an empty summary response. Please try again.")

        # Step 4: Save to Database
        summary = LectureSummary.objects.create(
            teacher=teacher,
            source_file_name=document.original_filename,
            summary_text=summary_text.strip(),
            is_satisfactory=True
        )

        return summary

    def export_pdf(self, summary_text: str, file_name: str = "Lecture_Summary") -> bytes:
        """Generates a clean PDF document from summary markdown/text using reportlab."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'SummaryTitle',
            parent=styles['Heading1'],
            fontSize=18,
            leading=22,
            spaceAfter=12,
            textColor='#1e293b'
        )

        body_style = ParagraphStyle(
            'SummaryBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=8,
            textColor='#334155'
        )

        story = [
            Paragraph(f"<b>AI Generated Lecture Summary: {file_name}</b>", title_style),
            Spacer(1, 12)
        ]

        lines = summary_text.splitlines()
        for line in lines:
            line_clean = line.strip()
            if not line_clean:
                story.append(Spacer(1, 6))
                continue
            
            # Simple escape for ReportLab XML tags
            safe_line = line_clean.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if safe_line.startswith("# "):
                story.append(Paragraph(f"<b>{safe_line[2:]}</b>", title_style))
            elif safe_line.startswith("## "):
                story.append(Paragraph(f"<b>{safe_line[3:]}</b>", styles['Heading2']))
            elif safe_line.startswith("- ") or safe_line.startswith("* "):
                story.append(Paragraph(f"• {safe_line[2:]}", body_style))
            else:
                story.append(Paragraph(safe_line, body_style))

        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes

    def export_docx(self, summary_text: str, file_name: str = "Lecture_Summary") -> bytes:
        """Generates a clean Word DOCX document from summary text using python-docx."""
        doc = DocxDocument()
        doc.add_heading(f"AI Generated Lecture Summary: {file_name}", level=1)
        
        lines = summary_text.splitlines()
        for line in lines:
            line_clean = line.strip()
            if not line_clean:
                continue
            if line_clean.startswith("# "):
                doc.add_heading(line_clean[2:], level=1)
            elif line_clean.startswith("## "):
                doc.add_heading(line_clean[3:], level=2)
            elif line_clean.startswith("### "):
                doc.add_heading(line_clean[4:], level=3)
            elif line_clean.startswith("- ") or line_clean.startswith("* "):
                doc.add_paragraph(line_clean[2:], style='List Bullet')
            else:
                doc.add_paragraph(line_clean)

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes
